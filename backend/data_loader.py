# data_loader.py

import os
import json
import fitz
import pandas as pd
from docx import Document
from pathlib import Path
from typing import List, Dict

class DocumentLoader:
    def _split_text_into_chunks(self, text: str, chunk_size: int = 500) -> List[str]:
        words = text.split()
        chunks, current_chunk, current_length = [], [], 0
        for word in words:
            current_chunk.append(word)
            current_length += len(word) + 1
            if current_length >= chunk_size:
                chunk_text = " ".join(current_chunk)
                if len(chunk_text.strip()) > 50:
                    chunks.append(chunk_text)
                current_chunk, current_length = [], 0
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            if len(chunk_text.strip()) > 50:
                chunks.append(chunk_text)
        return chunks

    def load(self, file_path: str) -> List[Dict]:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext == ".docx":
            return self._extract_docx(file_path)
        elif ext == ".txt":
            return self._extract_txt(file_path)
        elif ext == ".csv":
            return self._extract_csv(file_path)
        elif ext == ".json":
            return self._extract_json(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _extract_pdf(self, path: str) -> List[Dict]:
        doc = fitz.open(path)
        chunks = []
        for page_num in range(len(doc)):
            text = doc.load_page(page_num).get_text()
            for i, chunk in enumerate(self._split_text_into_chunks(text)):
                chunks.append({
                    "text": chunk,
                    "source": os.path.basename(path),
                    "file_type": "pdf",
                    "page": page_num + 1,
                    "chunk_id": len(chunks)
                })
        return chunks

    def _extract_docx(self, path: str) -> List[Dict]:
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                texts.append(row_text)
        return [{
            "text": chunk,
            "source": os.path.basename(path),
            "file_type": "docx",
            "chunk_id": i
        } for i, chunk in enumerate(self._split_text_into_chunks("\n".join(texts)))]

    def _extract_txt(self, path: str) -> List[Dict]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return [{
            "text": chunk,
            "source": os.path.basename(path),
            "file_type": "txt",
            "chunk_id": i
        } for i, chunk in enumerate(self._split_text_into_chunks(content))]

    def _extract_csv(self, path: str) -> List[Dict]:
        df = pd.read_csv(path)
        chunks = []
        for idx, row in df.iterrows():
            row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
            chunks.append({
                "text": row_text,
                "source": os.path.basename(path),
                "file_type": "csv",
                "row_number": idx + 1,
                "chunk_id": len(chunks)
            })
        return chunks

    def _extract_json(self, path: str) -> List[Dict]:
        def extract(obj, path=""):
            if isinstance(obj, dict):
                for k, v in obj.items():
                    yield from extract(v, f"{path}.{k}" if path else k)
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    yield from extract(item, f"{path}[{i}]")
            else:
                yield f"{path}: {obj}"
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return [{
            "text": chunk,
            "source": os.path.basename(path),
            "file_type": "json",
            "json_path": path,
            "chunk_id": i
        } for i, chunk in enumerate(self._split_text_into_chunks("\n".join(extract(data))))]

